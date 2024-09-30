import streamlit as st
import spacy
import re
import openai
import json
from docx import Document
from PyPDF2 import PdfReader
import os
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil
import docx
from docx.enum.table import WD_ROW_HEIGHT_RULE
import tempfile
from io import BytesIO
from dotenv import load_dotenv



load_dotenv()
OPENAI_API_KEY= os.getenv("OPENAI_API_KEY")
# print(OPENAI_API_KEY)

GPTModelLight = "gpt-4o-mini"
GPTModel = "gpt-4o"

def cleaner(text):
    clean_text = text.replace('*', '').replace('#', '')
    return clean_text
def header(doc_path, image):
    document = Document(doc_path)

    # Add header with table
    header = document.sections[0].header
    header.top_margin = Inches(1)

    htable = header.add_table(1, 2, Inches(8))
    htab_cells = htable.rows[0].cells
    
    row = htable.rows[0]
    row.height = Inches(1)  # Set the row height to 1 inch
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY 

    # Add image to the first cell
    ht0 = htab_cells[0].add_paragraph()
    kh = ht0.add_run()
    kh.add_picture(f'{image}', width=Inches(0.5), height=Inches(0.6))
    # Add styled text to the second cell
    ht1 = htab_cells[1].add_paragraph()
    ht1_run = ht1.add_run('CITY OF SANTA CRUZ, CALIFORNIA\nREQUEST FOR PROPOSALS – RFP NO. FN-0125 – DISASTER COST RECOVERY')
    ht1.alignment = WD_ALIGN_PARAGRAPH.RIGHT


    # Set font properties for the run
    font = ht1_run.font
    font.name = 'Calibri'
    font.size = Pt(8)
    font.bold = True
    document.save(doc_path)

def writer(doc_path, title, content):
    doc = Document(doc_path)
    section_found = False
    
    for i, paragraph in enumerate(doc.paragraphs):
        if title in paragraph.text:
            section_found = True
            insert_index = i + 1  # Insert after the title paragraph

            if i != 0:  # Avoid adding a page break at the beginning
                doc.paragraphs[insert_index].insert_paragraph_before().add_run().add_break()

            if isinstance(content, str):
                new_paragraph = doc.paragraphs[insert_index].insert_paragraph_before(content)
                new_paragraph.style = 'Normal'
                new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            elif isinstance(content, list):
                for item in content:
                    if item[0] == 'text':
                        new_paragraph = doc.paragraphs[insert_index].insert_paragraph_before(item[1])
                        new_paragraph.style = 'Normal'
                        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        insert_index += 1
                    
                    elif item[0] == 'table':
                        headers, rows = item[1], item[2]
                        table = doc.add_table(rows=1, cols=len(headers))
                        table.style = 'Table Grid'
                        
                        # Add headers
                        hdr_cells = table.rows[0].cells
                        for j, header in enumerate(headers):
                            hdr_cells[j].text = header
                        
                        # Add rows
                        for row in rows:
                            row_cells = table.add_row().cells
                            for j, cell_value in enumerate(row):
                                row_cells[j].text = str(cell_value)
                        
                        # Insert the table at the correct position
                        if insert_index < len(doc.paragraphs):
                            doc.paragraphs[insert_index].insert_paragraph_before()
                            table_paragraph = doc.paragraphs[insert_index]
                            table_paragraph._p.addnext(table._tbl)
                        else:
                            doc.add_paragraph()._p.addnext(table._tbl)
                        
                        insert_index += 1  # Account for the new paragraph and table
            break

    if not section_found:
        print(f"Section '{title}' not found in the document.")

    doc.save(doc_path)
    
def extract_text_from_pdf(uploaded_file):
    pdf_text = "" 
    try:
        reader = PdfReader(uploaded_file)
        num_pages = len(reader.pages)
        for page_num in range(num_pages):
            page = reader.pages[page_num]
            pdf_text += page.extract_text()
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
    return pdf_text
def read_word(file):
    doc = Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# CL_samples= read_word('files\\Cover letters samples.docx')
CL_samples = read_word(os.path.join('files', 'Cover letters samples.docx'))
# expertise= read_word('files\\expertise.docx')
expertise = read_word(os.path.join('files', 'expertise.docx'))
# team_qual= read_word('files\\team_qualification.docx')
team_qual = read_word(os.path.join('files', 'team_qualification.docx'))
# project_und= read_word('files\\project_understanding.docx')
project_und = read_word(os.path.join('files', 'project_understanding.docx'))
# technical_app= read_word('files\\technical_approach.docx')
technical_app = read_word(os.path.join('files', 'technical_approach.docx'))
# workplan= read_word('files\\work_plan.docx')
workplan = read_word(os.path.join('files', 'work_plan.docx'))
# cost_det= read_word('files\\cost_details.docx')
cost_det = read_word(os.path.join('files', 'cost_details.docx'))


from docx.shared import Inches

def create_docx_table(data, section):
    content_items = []
    
    if not isinstance(data, dict):
        raise TypeError(f"Expected 'data' to be a dictionary, but got {type(data).__name__}.")
    
    section_data = data[section]
    
    if section_data is None:
        raise ValueError(f"Section '{section}' not found in the data.")
 
    if 'intro' in section_data:
        content_items.append(('text', section_data['intro']))
    
    def create_table(table_data):
        if isinstance(table_data, dict) and 'headers' in table_data and 'rows' in table_data:
            headers = table_data['headers']
            rows = table_data['rows']
            return ('table', headers, rows)
        return None
    
    for key in ['table1', 'table2', 'table']:
        if key in section_data:
            table = create_table(section_data[key])
            if table:
                content_items.append(table)
    
    if 'conclusion' in section_data:
        content_items.append(('text', section_data['conclusion']))
    
    return content_items

def cover_letter(dets, scope):
     # Define the context for the summary
    context = (f'''make a relevant cover letter inspired by the one I did on my previous proposal
[CL_samples] 
and adapt it to this RFP. Here is the information for the purchasing manager, the person in charge and his information : {dets}
and here the scope of service of the RFP :
{scope}
It needs to end with these sentences : 
For any inquiries or additional information, please feel free to contact me directly at (518) 330-8526 or via email at jeff.abraham@grantcityllc.com We are excited about the prospect of becoming your trusted partner and advocate in this crucial endeavor.
Sincerely,
''')

    # Call the OpenAI API to generate a summary
    response =  openai.chat.completions.create(
        model=GPTModelLight,
        temperature=0.0,
        max_tokens=10000,
        messages=[
            {"role": "system", "content": context},
            {"role": "user", "content": CL_samples}
        ]
    )
    
    return response.choices[0].message.content.strip()
def FIRM_QUALIFICATIONS( scope):
     # Define the context for the summary
    context = (f'''Here is the section : FIRM QUALIFICATIONS, EXPERTISE, AND EXPERIENCE from a past proposal : 
 [expertise.docx] 
Adapt to this new RFP, here is the scope of service of the RFP : 
{scope}
 

''')

    # Call the OpenAI API to generate a summary
    response =  openai.chat.completions.create(
        model=GPTModelLight,
        temperature=0.0,
        max_tokens=10000,
        messages=[
            {"role": "system", "content": context},
            {"role": "user", "content": expertise}
        ]
    )
    
    return response.choices[0].message.content.strip()
def TEAM_QUALIFICATIONS(scope):
     # Define the context for the summary
    context = (f'''Here is the section “TEAM QUALIFICATIONS, EXPERTISE, AND EXPERIENCE” from a past proposal : 
 [team_qual] 
It needs to be assertive and persuasive, keep the same tone. Adapt to this new RFP , here is the scope of service of the RFP : 
{scope}

 

''')

    # Call the OpenAI API to generate a summary
    response =  openai.chat.completions.create(
        model=GPTModelLight,
        temperature=0.0,
        max_tokens=10000,
        messages=[
            {"role": "system", "content": context},
            {"role": "user", "content": team_qual}
        ]
    )
    
    return response.choices[0].message.content.strip()
def Project_understanding(scope, article):
     # Define the context for the summary
    context = (f'''write me a project understanding section for RFP, I am writing a proposal with a section project understanding , here is this section from a past proposal : [project_und] 
For this new RFP, here is the scope of service / work : {scope} 
Very important : our company named grant city LLC, adapt the text mentioning that we know all about the disaster event and is ready willing and able to assist (as needed). We add a recent press article about the disaster event, keep the same tone we used in the past proposal , project understanding section. Here is the article : {article}


''')

    # Call the OpenAI API to generate a summary
    response =  openai.chat.completions.create(
        model=GPTModelLight,
        temperature=0.0,
        max_tokens=10000,
        messages=[
            {"role": "system", "content": context},
            {"role": "user", "content": project_und}
        ]
    )
    
    return response.choices[0].message.content.strip()

def Technical_approach( scope):
     # Define the context for the summary
    context = (f'''Create a technical approach section for responding to an RFP in a JSON format.
        Detailed methodology to achieve the project's goals, identification of deliverables, key decision points, and any proposed changes to the scope of services.
        If you output a table, output table markdown formatted, any row in its own line. 
        Here is the Technical Approach from a past proposal : 
        [technical_app] 
        Adapt to this new RFP, Approach why our company : grant city LLC about managing the project is the best choice, it needs to be assertive and persuasive . Here is the scope of service of the RFP : 
        [{scope}].
        
        Return Output in JSON format:
        \u007b
            technical_approach : \u007b
                "intro": "make few paragraphs and points for intro taking reference example from [technical_app] provided to you.",
                "table": [
                    "headers": ["Phase", "Objective", "Deliverables", "Key Decision Points"],
                    "rows": [
                        ["Phase 1", "Objective", "Deliverables", "Key Decision Points"],
                        ["Phase 2", "Objective", "Deliverables", "Key Decision Points"],
                            .
                            .
                        //...more rows as needed
                    ]
                ],
                "conclusion": "make few paragraphs for conclusion taking reference example from [technical_app] provided to you."
            \u007b
        \u007b
    ''')

    # Call the OpenAI API to generate a summary
    response =  openai.chat.completions.create(
        model=GPTModelLight,
        temperature=0.0,
        max_tokens=10000,
        messages=[
            {"role": "system", "content": context},
            {"role": "user", "content": technical_app}
        ]
    )
    
    response_content = response.choices[0].message.content
    cleaned_response = response_content.replace('```json\n', '').replace('\n```', '')

    # Parse the JSON response
    try:
        structured_data = json.loads(cleaned_response)
    except json.JSONDecodeError as e:
        print("Failed to decode JSON:", e)
        return None
    return structured_data
 


def Work_plan(scope): 
    
     #extracting existing teams for proposed title and role.
    path = os.path.join('files', 'work_plan.docx')
    doc = Document(path)
    staff_titles_dict = {}
    
    for table in doc.tables:
        for row in table.rows[1:]:
            name = row.cells[0].text.strip()
            title = row.cells[1].text.strip()
            staff_titles_dict[name] = title
    
    names_and_titles = [f"{name}: {title}" for name, title in staff_titles_dict.items()]
    
     # Define the context for the summary 
    context = (f'''Create a work plan section for responding to an RFP in JSON format. 
    The work plan should include a detailed project schedule, a matrix of personnel tasks, and estimated effort in hours. 
    **Use the existing team members (names, titles, and roles) from the previous project as provided here {names_and_titles}. Do not create dummy names or new data for the team. Only adapt their specific tasks, duties, and hours to reflect the new RFP's scope of services.**

    Here is the work plan from a past proposal:
        [workplan]

    CRITICAL:**Ensure that the team members' names, titles, and roles remain unchanged. However, the specific duties, tasks, and hours allocated may change according to the new RFP requirements.**

    Here is the scope of service of the RFP: 
        [{scope}]

    Remember: The output should be similar in length and content to the past proposal, but adapt to the new RFP's requirements, maintaining a similar matrix structure.

    Output Format:
    \u007b
        work_plan : \u007b
            "intro": "Take reference example from [workplan] in length and tone provided to you for intro.",
            "table": [
                "headers": ["Proposed Staff", "Proposed Title", "Proposed Duty Statement", "SPECIFIC SOW ITEM PERFORMED BY RELEVANT STAFF(AS REQUESTED IN THE RFP)","ANNUAL ALLOCATION OF HOURS PER TASK"],
                "rows": [
                    ["Name of existing staff", "Title", "Duty", "Sow item by relevant staff","Annual Hours allocation"],
                    ["Name of existing staff", "Title", "Duty", "Sow item by relevant staff","Annual Hours allocation"],
                        .
                        .
                    ["Name of existing staff", "Title", "Duty", "Sow item by relevant staff","Annual Hours allocation"],
                ]
            ],
            "conclusion": "Take reference example from [workplan] in length and tone provided to you for conclusion."
        \u007b
    \u007b
    ''')

    # Call the OpenAI API to generate a summary
    response =  openai.chat.completions.create(
        model=GPTModelLight,
        temperature=0.0,
        max_tokens=10000,
        messages=[
            {"role": "system", "content": context},
            {"role": "user", "content": workplan}
        ]
    )
    
    response_content = response.choices[0].message.content
    cleaned_response = response_content.replace('```json\n', '').replace('\n```', '')

    # Parse the JSON response
    try:
        structured_data = json.loads(cleaned_response)
    except json.JSONDecodeError as e:
        print("Failed to decode JSON:", e)
        return None
    return structured_data

def COST_PROPOSAL(scope, budget):
     #extracting existing teams for proposed title and role.
    path = os.path.join('files', 'cost_details.docx')
    doc = Document(path)
    staff_titles_dict = {}
    
    for table in doc.tables:
        for row in table.rows[1:]:
            name = row.cells[0].text.strip()
            title = row.cells[1].text.strip()
            staff_titles_dict[name] = title
    
    names_and_titles = [f"{name}: {title}" for name, title in staff_titles_dict.items()]
         
     # Define the context for the summary
    context = (f'''Create a cost proposal section for responding to an RFP in JSON format. 
               
        CRITICAL:**Use the same team data/existing (names, titles, and roles) as provided here {names_and_titles} without generating new names or changing the team structure provided to you in [cost_det] ""
        Adapt the cost details, including rates, hours, and other expenses, to fit the new yearly budget and scope of services.

        Here is the cost proposal from a previous proposal we answered:
        [cost_det]

        Ensure the cost proposal reflects the existing team, but adjust their specific cost details to the new budget:
        {budget}.

        Here is the scope of service of the RFP:
        [{scope}]
        Output Format:
        \u007b
            "cost_proposal" : \u007b
                "intro": "Take reference example from [cost_det] in length and tone provided to you for intro.",
                "table1": [
                    "headers": ["Proposed Personnel", "Position Title", "Direct Labor Rate (Based on a normal 8-hour, 40 hours per week schedule)", "Indirect Labor Costs ","Professional Fee (Profit)","Other Direct Costs (ODC)","Fully Burdened Hourly Rate"],
                    "rows": [
                        ["Name of personnel", "Position", "Rate", "Indirect Labour cost","Profit","ODC","Hourly Rate"],
                        ["Name of personnel", "Position", "Rate", "Indirect Labour cost","Profit","ODC","Hourly Rate"],
                            .
                            .
                        ["Name of personnel", "Position", "Rate", "Indirect Labour cost","Profit","ODC","Hourly Rate"],
                    ]
                "table2": [
                    "headers": ["Initial Contract Period","Total Annual Not to Exceed Amount"],
                    "rows": [
                        ["Year", "Amount"]
                    ]
                ],
                "conclusion": "Take reference example from [cost_det] and provide exception and assumption for conclusion."
            \u007b
        \u007b   
    ''')

    # Call the OpenAI API to generate a summary
    response =  openai.chat.completions.create(
        model=GPTModelLight,
        temperature=0.0,
        max_tokens=10000,
        messages=[
            {"role": "system", "content": context},
            {"role": "user", "content": cost_det}
        ]
    )
    
    response_content = response.choices[0].message.content
    cleaned_response = response_content.replace('```json\n', '').replace('\n```', '')

    # Parse the JSON response
    try:
        structured_data = json.loads(cleaned_response)
    except json.JSONDecodeError as e:
        print("Failed to decode JSON:", e)
        return None
    return structured_data

def ref(scope):
     # Define the context for the summary
    context = (f'''I am creating a proposal to answer an RFP, my company GCC (Grant City LLC), we are writing the client reference section. I will send you an example of a past answer from a past proposal and you will adapt to the new RFP, with the new client name etc…
Here is a text from a past proposal for a previous client : 
GCC has provided descriptions and client references for projects of similar scope and size to the City’s project performed within the past three years. The City of Santy Cruz may contact any one of our client references during the evaluation phase. GCC manages more than 4,000 projects each year and has an extensive list of satisfied clients and award-winning performance on a wide range of diverse engineering, program management, and construction management assignments, including disaster recovery projects across the County. Figure 6 and Figure 7 provide a summary of our references and their similarity in scope, size, and complexity to Santa Cruz.
And here is the scope of service of the RFP : 
{scope}.



''')

    # Call the OpenAI API to generate a summary
    response =  openai.chat.completions.create(
        model=GPTModelLight,
        temperature=0.0,
        max_tokens=10000,
        messages=[
            {"role": "system", "content": context},
        ]
    )
    
    return response.choices[0].message.content.strip()


# Load the spaCy model
nlp = spacy.load("en_core_web_sm")

def extract_contact_info(text):
    # Regex patterns for email and phone number
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{1,4}?\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}'

    # Find all emails and phone numbers
    emails = re.findall(email_pattern, text)
    phones = re.findall(phone_pattern, text)

    return emails, phones

def main():
    st.title("Proposal Writer")
    
    # Add a space between the title and the subheader
    st.write("")  
    st.write("")  
    
    # File upload for RFP document
    st.subheader("Upload Documentation")
    uploaded_rfp = st.file_uploader("Files allowed: PDF or DOCX", type=["pdf", "docx"])
    
    purchasing_manager_details = []  
    scope_of_work = ""  

    if uploaded_rfp is not None:
        if uploaded_rfp.type == "application/pdf":
            rfp_text = extract_text_from_pdf(uploaded_rfp)
        elif uploaded_rfp.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            rfp_text = read_word(uploaded_rfp)
        
        # Automatically populate fields based on the extracted text
        if rfp_text:
            # Process the text with spaCy
            doc = nlp(rfp_text)

            # Extract contact info
            emails, phones = extract_contact_info(rfp_text)

            # Define keywords to search for purchasing manager info
            keywords = [
                "purchasing manager", 
                "procurement manager", 
                "procurement coordinator"
                "contact information",
                "department"
            ]
            
            for keyword in keywords:
                start_index = 0  
                while True:
                    start_index = rfp_text.lower().find(keyword, start_index)
                    if start_index == -1:
                        break  
                    
                    # Capture a larger context around the found keyword
                    context_start = max(0, start_index - 15)   
                    context_end = min(len(rfp_text), start_index + len(keyword) + 200) 
                    
                    # Extract and clean the context
                    context = rfp_text[context_start:context_end].strip()
                    
                    # Check if the context contains email, phone, and address
                    if any(email in context for email in emails) and \
                       any(phone in context for phone in phones):
                        purchasing_manager_details.append(context)
                    
                    # Move the start index forward to find the next instance
                    start_index += len(keyword)

            # Filter out irrelevant information
            purchasing_manager_details = [detail for detail in purchasing_manager_details if "do not contact" not in detail.lower()]
            
            # Store in session state
            st.session_state.purchasing_manager_details = "\n\n".join(purchasing_manager_details)  # Join instances with new lines
    
    # Purchasing Manager Details Section
    if uploaded_rfp is not None:
        st.subheader("Purchasing Manager Details")
        if 'purchasing_manager_details' in st.session_state:
            purchasing_manager_details = st.session_state.purchasing_manager_details
        st.text_area("", value=purchasing_manager_details, height=150)
        
        # Success message for RFP extraction
        if rfp_text:
            st.success("Purchasing Manager extracted successfully.")
        else:
            st.warning("Purchasing Manager not found in the RFP text.")
    
    # Scope of work / service Section
    if uploaded_rfp is not None:
        st.subheader("Scope of Work / Service")
        st.text_area("Extracted Text Below:", value=rfp_text, height=300)

        # Success message for RFP extraction
        if rfp_text:
            st.success("Scope of work / service extracted successfully.")
        else:
            st.warning("Scope of work / service not found in the RFP text.")



    # Article About Disaster Details
    st.subheader("Article About The Disaster")
    article_disaster = st.text_area("Article about the disaster", height=150)

        

   # Client Logo Upload
    st.subheader("Client Logo")
    client_logo = st.file_uploader("Upload your logo here", type=["png", "jpg", "jpeg"])
    if client_logo is not None:
        try:
            file_path = os.path.join("files", client_logo.name)
            with open(file_path, "wb") as f:
                f.write(client_logo.getbuffer())
            # relative_path= f"files\{client_logo.name}"
            relative_path = os.path.join("files", client_logo.name)  #this is for both linux and windows
        except Exception:
            st.error("File not uploaded properly.")
    else:
        st.warning("No logo uploaded. Proceeding without a client logo ?")
        relative_path = None  # set default logo path if required #  
        
    
        
    # Yearly Budget Input    
    st.subheader("Yearly Budget")
    yearly_budget = st.text_input("Yearly Budget")
    #uploaded_file = st.file_uploader("Upload a RFP", type="pdf") 


    # Other UI components remain unchanged...  
    if st.button("Generate Proposal"):
        # Generate content
        cl1 = cover_letter(purchasing_manager_details, scope_of_work)
        fq1 = FIRM_QUALIFICATIONS(scope_of_work)
        tq1 = TEAM_QUALIFICATIONS(scope_of_work)
        pu1 = Project_understanding(scope_of_work, article_disaster)
        ta1 = Technical_approach(scope_of_work)
        wp1 = Work_plan(scope_of_work)
        cp1 = COST_PROPOSAL(scope_of_work, yearly_budget)
        r1 = ref(scope_of_work)


        # Create tables
        ta1new = create_docx_table(ta1, 'technical_approach')
        wp1new = create_docx_table(wp1, 'work_plan')
        cp1new = create_docx_table(cp1, 'cost_proposal')

        # Clean content
        cl = cleaner(cl1)
        fq = cleaner(fq1)
        tq = cleaner(tq1)
        pu = cleaner(pu1)
        r = cleaner(r1)

        section5_content = [('text', "Project Understanding:\n" + pu + "\n\nTechnical Approach:")]
        section5_content.extend(ta1new)
        section5_content.append(('text', "\nWork Plan:"))
        section5_content.extend(wp1new)
        
        # Paths
        template_path = "template_proposal.docx"
        output_path = "proposal.docx"

        # Copy template
        shutil.copy(template_path, output_path)
        header(output_path, relative_path)
 
        # Modifications
        modifications = [
            ("COVER LETTER", cl),
            ("SECTION 3: FIRM QUALIFICATIONS, EXPERTISE, AND EXPERIENCE ", fq),
            ("SECTION 4: TEAM QUALIFICATIONS, EXPERTISE, AND EXPERIENCE ", tq),
            ("SECTION 5: TECHNICAL APPROACH AND WORK PLAN", ta1new),
            ("SECTION 5: TECHNICAL APPROACH AND WORK PLAN", wp1new),
            ("SECTION 5: TECHNICAL APPROACH AND WORK PLAN", pu),
            ("SECTION 6: COST PROPOSAL", cp1new),
            ("SECTION 7: REFERENCES ", r),
        ] 

 
        # Apply modifications to the document
        for title, content in modifications:
            print(f"Applying modification for {title}")
            writer(output_path, title, content)

        # Prepare for download
        output = BytesIO()
        op = Document(output_path)
        op.save(output)
        output.seek(0)
        st.download_button(
            label="Download Word File",
            data=output,
            file_name="proposal.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        os.remove(output_path)
        os.remove(relative_path)

if __name__ == "__main__":
    main()