import os
from docx import Document
from PyPDF2 import PdfReader
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE

def cleaner(text):
    clean_text = text.replace('*', '').replace('#', '')
    return clean_text


def header(doc_path, image):
    document = Document(doc_path)

    # Add header with table
    header = document.sections[0].header
    header.top_margin = Inches(1)

    htable = header.add_table(1, 2, Inches(8))
    htab_cells = htable.rows[0].cells
    
    row = htable.rows[0]
    row.height = Inches(1)  # Set the row height to 1 inch
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY 

    # Add image to the first cell
    ht0 = htab_cells[0].add_paragraph()
    kh = ht0.add_run()
    kh.add_picture(f'{image}', width=Inches(0.5), height=Inches(0.6))
    # Add styled text to the second cell
    ht1 = htab_cells[1].add_paragraph()
    ht1_run = ht1.add_run('CITY OF SANTA CRUZ, CALIFORNIA\nREQUEST FOR PROPOSALS – RFP NO. FN-0125 – DISASTER COST RECOVERY')
    ht1.alignment = WD_ALIGN_PARAGRAPH.RIGHT


    # Set font properties for the run
    font = ht1_run.font
    font.name = 'Calibri'
    font.size = Pt(8)
    font.bold = True
    document.save(doc_path)


def writer(doc_path, title, content):
    doc = Document(doc_path)
    section_found = False
    
    for i, paragraph in enumerate(doc.paragraphs):
        if title in paragraph.text:
            section_found = True
            insert_index = i + 1  # Insert after the title paragraph

            if i != 0:  # Avoid adding a page break at the beginning
                doc.paragraphs[insert_index].insert_paragraph_before().add_run().add_break()

            if isinstance(content, str):
                new_paragraph = doc.paragraphs[insert_index].insert_paragraph_before(content)
                new_paragraph.style = 'Normal'
                new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            elif isinstance(content, list):
                for item in content:
                    if item[0] == 'text':
                        new_paragraph = doc.paragraphs[insert_index].insert_paragraph_before(item[1])
                        new_paragraph.style = 'Normal'
                        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        insert_index += 1
                    
                    elif item[0] == 'table':
                        headers, rows = item[1], item[2]
                        table = doc.add_table(rows=1, cols=len(headers))
                        table.style = 'Table Grid'
                        
                        # Add headers
                        hdr_cells = table.rows[0].cells
                        for j, header in enumerate(headers):
                            hdr_cells[j].text = header
                        
                        # Add rows
                        for row in rows:
                            row_cells = table.add_row().cells
                            for j, cell_value in enumerate(row):
                                row_cells[j].text = str(cell_value)
                        
                        # Insert the table at the correct position
                        if insert_index < len(doc.paragraphs):
                            doc.paragraphs[insert_index].insert_paragraph_before()
                            table_paragraph = doc.paragraphs[insert_index]
                            table_paragraph._p.addnext(table._tbl)
                        else:
                            doc.add_paragraph()._p.addnext(table._tbl)
                        
                        insert_index += 1  # Account for the new paragraph and table
            break

    if not section_found:
        print(f"Section '{title}' not found in the document.")

    doc.save(doc_path)

    
def extract_text_from_pdf(uploaded_file):
    pdf_text = "" 
    try:
        reader = PdfReader(uploaded_file)
        num_pages = len(reader.pages)
        for page_num in range(num_pages):
            page = reader.pages[page_num]
            pdf_text += page.extract_text()
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
    return pdf_text


def read_word(file):
    doc = Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def create_docx_table(data, section):
    content_items = []
    
    if not isinstance(data, dict):
        raise TypeError(f"Expected 'data' to be a dictionary, but got {type(data).__name__}.")
    
    section_data = data[section]
    
    if section_data is None:
        raise ValueError(f"Section '{section}' not found in the data.")
 
    if 'intro' in section_data:
        content_items.append(('text', section_data['intro']))
    
    def create_table(table_data):
        if isinstance(table_data, dict) and 'headers' in table_data and 'rows' in table_data:
            headers = table_data['headers']
            rows = table_data['rows']
            return ('table', headers, rows)
        return None
    
    for key in ['table1', 'table2', 'table']:
        if key in section_data:
            table = create_table(section_data[key])
            if table:
                content_items.append(table)
    
    if 'conclusion' in section_data:
        content_items.append(('text', section_data['conclusion']))
    
    return content_items
